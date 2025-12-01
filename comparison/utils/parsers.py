import io
from typing import Optional

import pandas as pd
from docx import Document
from PyPDF2 import PdfReader


# Normalized column names we expect
NORMALIZED_COLUMNS = ["MPN", "Quantity", "Ref_Des", "Description"]


# -----------------------------
#  UNIVERSAL SAFE CELL CLEANER
# -----------------------------
def clean_cell(x):
    """
    Convert any messy cell (Series, list, NaN, tuple, dict, None)
    into a clean string or integer-friendly value.
    """
    try:
        # If cell is a Pandas Series → take first value
        if isinstance(x, pd.Series):
            x = x.iloc[0] if not x.empty else ""

        # If cell is a list or tuple → join
        if isinstance(x, (list, tuple)):
            return ", ".join(str(i) for i in x)

        # If cell is a dict → map keys/values
        if isinstance(x, dict):
            return ", ".join(f"{k}:{v}" for k, v in x.items())

        # If NULL/NaN
        if pd.isna(x):
            return ""

        # Everything else → string
        return str(x).strip()

    except Exception:
        return ""  # final fallback to avoid ANY crash


# -----------------------------------
#  MAIN NORMALIZATION FUNCTION
# -----------------------------------
def _normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    """
    Standardize and clean BOM data into columns:
    MPN, Quantity, Ref_Des, Description

    Handles:
    - duplicate headers
    - no headers
    - Series/list/dict cells
    - random PDF-extracted junk
    - numeric/float/NaN rows
    """
    if df is None or df.empty:
        return pd.DataFrame(columns=NORMALIZED_COLUMNS)

    df = df.copy()

    # Remove duplicate column names
    df = df.loc[:, ~df.columns.duplicated()]

    # Try intelligent header mapping
    col_map = {}
    for col in df.columns:
        col_lower = str(col).strip().lower()
        if "mpn" in col_lower or "part" in col_lower:
            col_map[col] = "MPN"
        elif "qty" in col_lower or "quantity" in col_lower:
            col_map[col] = "Quantity"
        elif "ref" in col_lower:
            col_map[col] = "Ref_Des"
        elif "desc" in col_lower or "description" in col_lower:
            col_map[col] = "Description"

    df = df.rename(columns=col_map)

    # Add missing required columns
    for column in NORMALIZED_COLUMNS:
        if column not in df.columns:
            df[column] = ""

    # Restrict to normalized columns only
    df = df[NORMALIZED_COLUMNS]

    # Clean every cell in every column
    for col in NORMALIZED_COLUMNS:
        df[col] = df[col].apply(clean_cell)

    # Convert Quantity to integer
    df["Quantity"] = (
        pd.to_numeric(df["Quantity"], errors="coerce")
        .fillna(0)
        .astype(int)
    )

    return df


# -----------------------------------
#  FILE PARSERS
# -----------------------------------
def parse_xlsx(django_file) -> pd.DataFrame:
    """
    Parse XLSX using pandas.
    """
    try:
        df = pd.read_excel(django_file)
        return _normalize_columns(df)
    except Exception:
        return pd.DataFrame(columns=NORMALIZED_COLUMNS)


def parse_csv(django_file) -> pd.DataFrame:
    """
    Parse CSV safely.
    """
    try:
        df = pd.read_csv(django_file)
        return _normalize_columns(df)
    except Exception:
        return pd.DataFrame(columns=NORMALIZED_COLUMNS)


def parse_txt(django_file, delimiter: Optional[str] = None) -> pd.DataFrame:
    """
    Parse TXT as whitespace or delimiter-separated values.
    """
    content = django_file.read()
    if isinstance(content, bytes):
        content = content.decode(errors="ignore")

    lines = [line.strip() for line in content.splitlines() if line.strip()]
    if not lines:
        return pd.DataFrame(columns=NORMALIZED_COLUMNS)

    rows = []
    for line in lines:
        parts = line.split(delimiter) if delimiter else line.split()
        rows.append(parts)

    df = pd.DataFrame(rows)
    return _normalize_columns(df)


def parse_docx(django_file) -> pd.DataFrame:
    """
    Parse DOCX table or paragraph text into DataFrame.
    """
    try:
        document = Document(django_file)

        if document.tables:
            table = document.tables[0]
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if len(rows) > 1:
                df = pd.DataFrame(rows[1:], columns=rows[0])
            else:
                df = pd.DataFrame(rows)
        else:
            lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            rows = [line.split() for line in lines]
            df = pd.DataFrame(rows)

        return _normalize_columns(df)

    except Exception:
        return pd.DataFrame(columns=NORMALIZED_COLUMNS)


def parse_pdf(django_file) -> pd.DataFrame:
    """
    Simple PDF parser.
    Works for text-based PDFs, not scanned images.
    """
    try:
        buffer = io.BytesIO(django_file.read())
        reader = PdfReader(buffer)

        text_lines = []
        for page in reader.pages:
            text = page.extract_text() or ""
            for line in text.splitlines():
                if line.strip():
                    text_lines.append(line.strip())

        rows = [line.split() for line in text_lines]
        df = pd.DataFrame(rows)
        return _normalize_columns(df)

    except Exception:
        return pd.DataFrame(columns=NORMALIZED_COLUMNS)


# -----------------------------------
#  DISPATCHER
# -----------------------------------
def parse_file_to_df(uploaded_file_obj, file_type: str) -> pd.DataFrame:
    """
    Decide which parser to use based on file extension.
    """
    file_type = file_type.lower().strip()

    if file_type == "xlsx":
        return parse_xlsx(uploaded_file_obj)

    if file_type == "csv":
        return parse_csv(uploaded_file_obj)

    if file_type == "txt":
        return parse_txt(uploaded_file_obj)

    if file_type == "docx":
        return parse_docx(uploaded_file_obj)

    if file_type == "pdf":
        return parse_pdf(uploaded_file_obj)

    # Unknown → return empty normalized DF
    return pd.DataFrame(columns=NORMALIZED_COLUMNS)
